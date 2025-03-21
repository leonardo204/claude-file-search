"""
Word (.docx) 파일 핸들러

이 모듈은 Word 문서에서 텍스트를 추출하는 기능을 제공합니다.
"""

import logging
import traceback
from pathlib import Path
from typing import List, Dict, Any

from file_handlers.base import FileHandler

logger = logging.getLogger("file_search.docx_handler")

class DOCXHandler(FileHandler):
    """Word 파일 핸들러"""
    
    def get_supported_extensions(self) -> List[str]:
        """
        지원하는 파일 확장자 목록을 반환합니다.
        
        Returns:
            지원하는 파일 확장자 목록
        """
        return ['.docx']
    
    def get_type_description(self) -> str:
        """
        파일 형식에 대한 설명을 반환합니다.
        
        Returns:
            파일 형식 설명
        """
        return 'Word'
    
    def extract_text(self, file_path: Path, max_content_sections: int = 10) -> List[Dict[str, Any]]:
        """
        Word 파일에서 텍스트를 추출합니다.
        
        Args:
            file_path: 파일 경로
            max_content_sections: 처리할 최대 섹션 수 (여기서는 문단 기준으로 분리)
            
        Returns:
            문단별 텍스트 정보 리스트
        """
        try:
            import docx
            
            logger.info(f"Extracting text from Word file: {file_path}")
            
            doc = docx.Document(file_path)
            
            # 문단 추출
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if i >= max_content_sections:
                    paragraphs.append({
                        "section_number": i + 1,
                        "section_type": "paragraph",
                        "text": "... 추가 문단 있음 (처리되지 않음) ..."
                    })
                    break
                    
                text = para.text.strip()
                if text:  # 빈 문단 건너뛰기
                    paragraphs.append({
                        "section_number": i + 1,
                        "section_type": "paragraph",
                        "text": text
                    })
            
            # 표 추출 (각 표는 별도의 섹션으로 처리)
            tables_processed = 0
            table_section_number = len(paragraphs) + 1
            
            for table in doc.tables:
                if tables_processed >= max_content_sections // 2:  # 테이블은 최대 섹션의 절반까지만 처리
                    continue
                    
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        table_text += cell.text + "\t"
                    table_text += "\n"
                
                if table_text.strip():
                    paragraphs.append({
                        "section_number": table_section_number,
                        "section_type": "table",
                        "text": table_text.strip()
                    })
                    table_section_number += 1
                    tables_processed += 1
            
            return paragraphs
        except Exception as e:
            logger.error(f"Error extracting text from Word file {file_path}: {e}")
            logger.error(traceback.format_exc())
            return [{"section_number": 0, "section_type": "error", "text": f"Word 파일 처리 중 오류 발생: {str(e)}"}]